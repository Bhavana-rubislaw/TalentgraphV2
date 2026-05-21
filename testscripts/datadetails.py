"""
datadetails.py
==============================================================================
TalentGraph V2 - Full database data report.

Run from the testscripts folder (venv must be active):
    cd "..\backend2"
    .\venv\Scripts\Activate.ps1
    cd "..\testscripts"
    python datadetails.py

What it does:
  1. Prints a structured report to the console.
  2. Saves an identically structured Word document (.docx) to the same
     folder as this script, named:
         TalentGraph_Report_YYYY-MM-DD_HHMMSS.docx
     A fresh file is generated on every run so the history is preserved
     and the report always reflects the latest database state.

Sections covered:
  * Overall database summary
  * Per-candidate: profile, social links, certifications, resumes,
                   job preferences with skills and location preferences
  * Per-company  : profile info, job postings with required skills
  * Count tables A-F for quick manager reporting
  * Quick-stats snapshot
==============================================================================
"""

import sys
import os
import io
from datetime import datetime

# -- Windows terminal Unicode fix ---------------------------------------------
if hasattr(sys.stdout, 'buffer') and sys.stdout.encoding.lower() != 'utf-8':
    sys.stdout = io.TextIOWrapper(sys.stdout.buffer, encoding='utf-8', errors='replace')

# -- Locate backend2/app so imports work regardless of cwd -------------------
SCRIPT_DIR  = os.path.dirname(os.path.abspath(__file__))
BACKEND_DIR = os.path.abspath(os.path.join(SCRIPT_DIR, "..", "backend2"))
sys.path.insert(0, BACKEND_DIR)

from sqlmodel import Session, select
from app.database import engine
from app.models import (
    User, Candidate, Company,
    Resume, Certification,
    JobProfile, Skill, LocationPreference,
    JobPosting, JobPostingSkill,
    Application, Match, Swipe,
)

# -- python-docx --------------------------------------------------------------
try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False
    print("[WARN] python-docx not found - Word export disabled. "
          "Install with: pip install python-docx")

# =============================================================================
# DATA COLLECTION  (single DB round-trip shared by console + docx)
# =============================================================================

def collect_data(session):
    users        = session.exec(select(User).order_by(User.id)).all()
    candidates   = session.exec(select(Candidate).order_by(Candidate.id)).all()
    companies    = session.exec(select(Company).order_by(Company.id)).all()
    resumes      = session.exec(select(Resume)).all()
    certs        = session.exec(select(Certification)).all()
    profiles     = session.exec(select(JobProfile).order_by(JobProfile.id)).all()
    skills       = session.exec(select(Skill)).all()
    loc_prefs    = session.exec(select(LocationPreference)).all()
    postings     = session.exec(select(JobPosting).order_by(JobPosting.id)).all()
    post_skills  = session.exec(select(JobPostingSkill)).all()
    applications = session.exec(select(Application)).all()
    matches      = session.exec(select(Match)).all()
    swipes       = session.exec(select(Swipe)).all()

    user_by_id         = {u.id: u for u in users}
    resumes_by_cand    = {}
    for r in resumes:    resumes_by_cand.setdefault(r.candidate_id, []).append(r)
    certs_by_cand      = {}
    for c in certs:      certs_by_cand.setdefault(c.candidate_id, []).append(c)
    profiles_by_cand   = {}
    for p in profiles:   profiles_by_cand.setdefault(p.candidate_id, []).append(p)
    skills_by_profile  = {}
    for sk in skills:    skills_by_profile.setdefault(sk.job_profile_id, []).append(sk)
    locs_by_profile    = {}
    for lp in loc_prefs: locs_by_profile.setdefault(lp.job_profile_id, []).append(lp)
    postings_by_comp   = {}
    for p in postings:   postings_by_comp.setdefault(p.company_id, []).append(p)
    pskills_by_posting = {}
    for ps in post_skills: pskills_by_posting.setdefault(ps.job_posting_id, []).append(ps)
    apps_by_cand       = {}
    for a in applications: apps_by_cand.setdefault(a.candidate_id, []).append(a)
    matches_by_cand    = {}
    for m in matches:    matches_by_cand.setdefault(m.candidate_id, []).append(m)

    return dict(
        users=users, candidates=candidates, companies=companies,
        resumes=resumes, certs=certs, profiles=profiles,
        skills=skills, loc_prefs=loc_prefs, postings=postings,
        post_skills=post_skills, applications=applications,
        matches=matches, swipes=swipes,
        user_by_id=user_by_id,
        resumes_by_cand=resumes_by_cand, certs_by_cand=certs_by_cand,
        profiles_by_cand=profiles_by_cand, skills_by_profile=skills_by_profile,
        locs_by_profile=locs_by_profile, postings_by_comp=postings_by_comp,
        pskills_by_posting=pskills_by_posting,
        apps_by_cand=apps_by_cand, matches_by_cand=matches_by_cand,
    )

# =============================================================================
# CONSOLE REPORT
# =============================================================================

def _hdr(text, char="="):
    w = 90
    print(); print(char * w); print(f"  {text}"); print(char * w)

def _row(*cols, widths=None):
    if widths is None:
        print("  " + "  ".join(str(c) for c in cols))
    else:
        print("  " + "  ".join(str(c)[:w].ljust(w) for c, w in zip(cols, widths)))

def _tbl_hdr(*cols, widths=None):
    _row(*cols, widths=widths)
    if widths:
        print("  " + "  ".join("-" * w for w in widths))

def _v(x):
    return x.value if hasattr(x, 'value') else str(x)


def print_report(d):
    users=d["users"]; candidates=d["candidates"]; companies=d["companies"]
    resumes=d["resumes"]; certs=d["certs"]; profiles=d["profiles"]
    skills=d["skills"]; loc_prefs=d["loc_prefs"]
    postings=d["postings"]; post_skills=d["post_skills"]
    applications=d["applications"]; matches=d["matches"]; swipes=d["swipes"]
    user_by_id=d["user_by_id"]
    resumes_by_cand=d["resumes_by_cand"]; certs_by_cand=d["certs_by_cand"]
    profiles_by_cand=d["profiles_by_cand"]; skills_by_profile=d["skills_by_profile"]
    locs_by_profile=d["locs_by_profile"]; postings_by_comp=d["postings_by_comp"]
    pskills_by_posting=d["pskills_by_posting"]
    apps_by_cand=d["apps_by_cand"]; matches_by_cand=d["matches_by_cand"]

    total_cand_skills = sum(len(skills_by_profile.get(p.id,[])) for p in profiles)
    total_comp_skills = sum(len(pskills_by_posting.get(p.id,[])) for p in postings)

    # Overall summary
    _hdr("OVERALL DATABASE SUMMARY")
    W=[42,10]
    _tbl_hdr("Metric","Count",widths=W)
    _row("Total Users",len(users),widths=W)
    _row("  Candidates",len(candidates),widths=W)
    _row("  Companies (Recruiters/HR)",len(companies),widths=W)
    _row("  Other (admin/system)",len(users)-len(candidates)-len(companies),widths=W)
    print()
    _row("Total Resumes uploaded",len(resumes),widths=W)
    _row("Total Certifications",len(certs),widths=W)
    print()
    _row("Total Job Profiles (preferences)",len(profiles),widths=W)
    _row("  Skills across all profiles",total_cand_skills,widths=W)
    _row("  Location prefs across all",len(loc_prefs),widths=W)
    print()
    _row("Total Job Postings",len(postings),widths=W)
    _row("  Skills across all postings",total_comp_skills,widths=W)
    print()
    _row("Total Applications",len(applications),widths=W)
    _row("Total Matches",len(matches),widths=W)
    _row("Total Swipes",len(swipes),widths=W)

    # Candidate details
    _hdr("CANDIDATE DETAILS  (one block per candidate)")
    for cand in candidates:
        user=user_by_id.get(cand.user_id)
        c_res=resumes_by_cand.get(cand.id,[]); c_crt=certs_by_cand.get(cand.id,[])
        c_prof=profiles_by_cand.get(cand.id,[]); c_apps=apps_by_cand.get(cand.id,[])
        print("-"*90)
        print(f"  CANDIDATE ID: {cand.id}  |  User ID: {cand.user_id}")
        print(f"  Name   : {cand.name}")
        print(f"  Email  : {cand.email}  (login: {user.email if user else '?'})")
        print(f"  Phone  : {cand.phone}")
        print(f"  Location: {cand.residential_address}, {cand.location_state} {cand.location_zipcode}")
        socials=[]
        if cand.linkedin_url:  socials.append(f"LinkedIn: {cand.linkedin_url}")
        if cand.github_url:    socials.append(f"GitHub: {cand.github_url}")
        if cand.portfolio_url: socials.append(f"Portfolio: {cand.portfolio_url}")
        if socials:
            print(f"  Social Links ({len(socials)}):")
            for sl in socials: print(f"      * {sl}")
        else:
            print("  Social Links: none")
        print(f"  Resumes: {len(c_res)}  |  Certifications: {len(c_crt)}  |  Applications: {len(c_apps)}")
        for c in c_crt:
            print(f"      * [{c.id}] {c.name}  |  {c.issuer or 'N/A'}  |  Issued: {c.issued_date or 'N/A'}  Exp: {c.expiry_date or 'N/A'}")
        print(f"  Job Profiles: {len(c_prof)}")
        for jp in c_prof:
            jp_sk=skills_by_profile.get(jp.id,[]); jp_lc=locs_by_profile.get(jp.id,[])
            tech=[s for s in jp_sk if s.skill_category=="technical"]
            func=[s for s in jp_sk if s.skill_category=="functional"]
            soft=[s for s in jp_sk if s.skill_category=="soft"]
            print()
            print(f"      +-- [{jp.id}] {jp.profile_name}")
            print(f"      |  {jp.product_vendor or 'N/A'} / {jp.product_type or 'N/A'} / {jp.job_role or 'N/A'}")
            print(f"      |  {jp.years_of_experience} yrs  |  {jp.seniority_level or 'N/A'}  |  {_v(jp.worktype)}  |  {_v(jp.employment_type)}")
            print(f"      |  {_v(jp.salary_currency)} {jp.salary_min:,.0f} - {jp.salary_max:,.0f}  |  Visa: {_v(jp.visa_status)}")
            if jp_lc: print(f"      |  Locations: " + " | ".join(f"{l.city}, {l.state}" for l in jp_lc))
            print(f"      |  Skills: {len(jp_sk)} (Tech:{len(tech)} Func:{len(func)} Soft:{len(soft)})")
            if tech: print(f"      |    Tech : " + ", ".join(f"{s.skill_name}(L{s.proficiency_level})" for s in tech))
            if func: print(f"      |    Func : " + ", ".join(f"{s.skill_name}(L{s.proficiency_level})" for s in func))
            if soft: print(f"      |    Soft : " + ", ".join(f"{s.skill_name}(L{s.proficiency_level})" for s in soft))
            print(f"      +------------------------------------------------------")

    # Company details
    _hdr("COMPANY DETAILS  (one block per company/recruiter)")
    for comp in companies:
        user=user_by_id.get(comp.user_id); c_posts=postings_by_comp.get(comp.id,[])
        print("-"*90)
        print(f"  COMPANY ID: {comp.id}  |  {comp.company_name}")
        print(f"  Email: {comp.company_email}  |  Type: {comp.employee_type}  |  Dept: {comp.department or 'N/A'}")
        print(f"  Location: {comp.company_location or 'N/A'}  |  Phone: {comp.phone_number or 'N/A'}")
        print(f"  Website: {comp.company_website or 'N/A'}")
        if comp.linkedin_profile: print(f"  LinkedIn: {comp.linkedin_profile}")
        print(f"  Hiring Focus: {comp.hiring_focus or 'N/A'}  |  Profile Complete: {comp.profile_complete}")
        print(f"  Job Postings: {len(c_posts)}")
        for jp in c_posts:
            jp_sk=pskills_by_posting.get(jp.id,[]); tech=[s for s in jp_sk if s.skill_category=="technical"]; soft=[s for s in jp_sk if s.skill_category=="soft"]
            print()
            print(f"      +-- [{jp.id}] {jp.job_title}")
            print(f"      |  {jp.product_vendor} / {jp.product_type} / {jp.job_role}  |  {jp.seniority_level}")
            print(f"      |  {_v(jp.worktype)}  |  {_v(jp.employment_type)}  |  {jp.location}  |  Start: {jp.start_date}")
            print(f"      |  {_v(jp.salary_currency)} {jp.salary_min:,.0f} - {jp.salary_max:,.0f}  |  Status: {_v(jp.status)}")
            print(f"      |  Skills: {len(jp_sk)} (Tech:{len(tech)} Soft:{len(soft)})")
            if tech: print(f"      |    Tech : " + ", ".join(f"{s.skill_name}(R{s.rating})" for s in tech))
            if soft: print(f"      |    Soft : " + ", ".join(f"{s.skill_name}(R{s.rating})" for s in soft))
            print(f"      +------------------------------------------------------")

    # Count tables
    _hdr("COUNT SUMMARY TABLES")
    total_cand_sk = sum(len(skills_by_profile.get(p.id,[])) for p in profiles)
    total_comp_sk = sum(len(pskills_by_posting.get(p.id,[])) for p in postings)

    print("\n  TABLE A: Candidates - Data Completeness")
    W2=[4,28,7,6,8,10,6,10]
    _tbl_hdr("ID","Name","Resumes","Certs","Profiles","PrfSkills","Locs","Apps",widths=W2)
    for cand in candidates:
        cp=profiles_by_cand.get(cand.id,[])
        _row(cand.id,cand.name[:26],len(resumes_by_cand.get(cand.id,[])),
             len(certs_by_cand.get(cand.id,[])),len(cp),
             sum(len(skills_by_profile.get(p.id,[])) for p in cp),
             sum(len(locs_by_profile.get(p.id,[])) for p in cp),
             len(apps_by_cand.get(cand.id,[])),widths=W2)

    print("\n  TABLE B: Companies - Job Postings & Skills")
    W3=[4,26,16,10,14]
    _tbl_hdr("ID","Company","EmployeeType","Postings","PostingSkills",widths=W3)
    for comp in companies:
        cp=postings_by_comp.get(comp.id,[])
        _row(comp.id,comp.company_name[:24],comp.employee_type[:14],
             len(cp),sum(len(pskills_by_posting.get(p.id,[])) for p in cp),widths=W3)

    print("\n  TABLE C: Job Profiles - Skills Breakdown")
    W4=[4,4,38,9,10,8,7]
    _tbl_hdr("PID","CID","ProfileName","Technical","Functional","Soft","Total",widths=W4)
    for jp in profiles:
        jp_sk=skills_by_profile.get(jp.id,[])
        _row(jp.id,jp.candidate_id,jp.profile_name[:36],
             sum(1 for s in jp_sk if s.skill_category=="technical"),
             sum(1 for s in jp_sk if s.skill_category=="functional"),
             sum(1 for s in jp_sk if s.skill_category=="soft"),
             len(jp_sk),widths=W4)

    print("\n  TABLE D: Job Postings - Skills Breakdown")
    W5=[4,4,40,9,8,7]
    _tbl_hdr("JID","CID","JobTitle","Technical","Soft","Total",widths=W5)
    for jp in postings:
        jp_sk=pskills_by_posting.get(jp.id,[])
        _row(jp.id,jp.company_id,jp.job_title[:38],
             sum(1 for s in jp_sk if s.skill_category=="technical"),
             sum(1 for s in jp_sk if s.skill_category=="soft"),
             len(jp_sk),widths=W5)

    print("\n  TABLE E: Social Links per Candidate")
    W6=[4,26,10,10,12]
    _tbl_hdr("CID","Name","LinkedIn","GitHub","Portfolio",widths=W6)
    for cand in candidates:
        _row(cand.id,cand.name[:24],
             "Yes" if cand.linkedin_url else "No",
             "Yes" if cand.github_url else "No",
             "Yes" if cand.portfolio_url else "No",widths=W6)

    print("\n  TABLE F: Social Links per Company")
    W7=[4,26,12]
    _tbl_hdr("CID","Company","LinkedIn",widths=W7)
    for comp in companies:
        _row(comp.id,comp.company_name[:24],
             "Yes" if comp.linkedin_profile else "No",widths=W7)

    # Quick stats
    _hdr("QUICK STATS SNAPSHOT  (for manager reporting)", char="-")
    avg_pr=len(profiles)/max(len(candidates),1); avg_ct=len(certs)/max(len(candidates),1)
    avg_po=len(postings)/max(len(companies),1)
    avg_ps=total_cand_sk/max(len(profiles),1); avg_jp=total_comp_sk/max(len(postings),1)
    print(f"  {'Metric':<52} {'Value':>10}")
    print(f"  {'-'*62}")
    print(f"  {'Total users in system':<52} {len(users):>10}")
    print(f"  {'  - Candidates':<52} {len(candidates):>10}")
    print(f"  {'  - Companies':<52} {len(companies):>10}")
    print(f"  {'Avg job profiles per candidate':<52} {avg_pr:>10.2f}")
    print(f"  {'Avg certifications per candidate':<52} {avg_ct:>10.2f}")
    print(f"  {'Avg job postings per company':<52} {avg_po:>10.2f}")
    print(f"  {'Avg skills per job profile':<52} {avg_ps:>10.2f}")
    print(f"  {'Avg skills per job posting':<52} {avg_jp:>10.2f}")
    print(f"  {'Total applications in system':<52} {len(applications):>10}")
    print(f"  {'Total matches in system':<52} {len(matches):>10}")
    print(f"  {'Total swipes in system':<52} {len(swipes):>10}")
    _hdr("END OF REPORT")
    sys.stdout.flush()

# =============================================================================
# WORD DOCUMENT REPORT
# =============================================================================

def _set_cell_bg(cell, hex_color):
    tc=cell._tc; tcPr=tc.get_or_add_tcPr()
    shd=OxmlElement('w:shd')
    shd.set(qn('w:val'),'clear'); shd.set(qn('w:color'),'auto')
    shd.set(qn('w:fill'),hex_color); tcPr.append(shd)

def _add_table(doc, headers, rows, col_widths=None, header_bg="1F3864"):
    table=doc.add_table(rows=1+len(rows),cols=len(headers))
    table.style='Table Grid'
    hdr_cells=table.rows[0].cells
    for i,h in enumerate(headers):
        hdr_cells[i].text=h
        p=hdr_cells[i].paragraphs[0]; p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        run=p.runs[0] if p.runs else p.add_run(h)
        run.font.bold=True; run.font.color.rgb=RGBColor(0xFF,0xFF,0xFF); run.font.size=Pt(9)
        _set_cell_bg(hdr_cells[i],header_bg)
    for ri,row_data in enumerate(rows):
        row_cells=table.rows[ri+1].cells
        bg="EEF2FF" if ri%2==0 else "FFFFFF"
        for ci,val in enumerate(row_data):
            row_cells[ci].text=str(val)
            if row_cells[ci].paragraphs[0].runs:
                row_cells[ci].paragraphs[0].runs[0].font.size=Pt(9)
            _set_cell_bg(row_cells[ci],bg)
    if col_widths:
        for row in table.rows:
            for ci,w in enumerate(col_widths):
                if ci<len(row.cells): row.cells[ci].width=w

def _wdg_heading(doc, text, level=1):
    p=doc.add_heading(text,level=level)
    color=RGBColor(0x1F,0x38,0x64) if level==1 else RGBColor(0x2E,0x74,0xB5)
    if p.runs: p.runs[0].font.color.rgb=color

def generate_word_report(d, output_path):
    if not DOCX_AVAILABLE:
        print("[SKIP] python-docx not available.")
        return ""

    doc=Document()
    sec=doc.sections[0]
    sec.page_width=Inches(11); sec.page_height=Inches(8.5)
    sec.left_margin=sec.right_margin=Inches(0.75)
    sec.top_margin=sec.bottom_margin=Inches(0.75)

    # Title
    title=doc.add_heading("TalentGraph V2 - Database Report",0)
    title.alignment=WD_ALIGN_PARAGRAPH.CENTER
    if title.runs: title.runs[0].font.color.rgb=RGBColor(0x1F,0x38,0x64)
    sub=doc.add_paragraph(); sub.alignment=WD_ALIGN_PARAGRAPH.CENTER
    run=sub.add_run(f"Generated: {datetime.now().strftime('%B %d, %Y  %I:%M %p')}")
    run.font.size=Pt(11); run.font.italic=True
    run.font.color.rgb=RGBColor(0x44,0x44,0x44)
    doc.add_paragraph()

    users=d["users"]; candidates=d["candidates"]; companies=d["companies"]
    resumes=d["resumes"]; certs=d["certs"]; profiles=d["profiles"]
    skills=d["skills"]; loc_prefs=d["loc_prefs"]
    postings=d["postings"]; post_skills=d["post_skills"]
    applications=d["applications"]; matches=d["matches"]; swipes=d["swipes"]
    user_by_id=d["user_by_id"]
    resumes_by_cand=d["resumes_by_cand"]; certs_by_cand=d["certs_by_cand"]
    profiles_by_cand=d["profiles_by_cand"]; skills_by_profile=d["skills_by_profile"]
    locs_by_profile=d["locs_by_profile"]; postings_by_comp=d["postings_by_comp"]
    pskills_by_posting=d["pskills_by_posting"]
    apps_by_cand=d["apps_by_cand"]; matches_by_cand=d["matches_by_cand"]

    total_cand_sk=sum(len(skills_by_profile.get(p.id,[])) for p in profiles)
    total_comp_sk=sum(len(pskills_by_posting.get(p.id,[])) for p in postings)
    avg_pr=len(profiles)/max(len(candidates),1); avg_ct=len(certs)/max(len(candidates),1)
    avg_po=len(postings)/max(len(companies),1)
    avg_ps=total_cand_sk/max(len(profiles),1); avg_jp=total_comp_sk/max(len(postings),1)

    # ----- Section 1: Overall Summary ----------------------------------------
    _wdg_heading(doc,"1. Overall Database Summary")
    _add_table(doc,["Metric","Count / Value"],[
        ["Total Users",str(len(users))],
        ["  Candidates",str(len(candidates))],
        ["  Companies (Recruiters / HR)",str(len(companies))],
        ["  Other (admin / system)",str(len(users)-len(candidates)-len(companies))],
        ["Total Resumes Uploaded",str(len(resumes))],
        ["Total Certifications",str(len(certs))],
        ["Total Job Profiles (preferences)",str(len(profiles))],
        ["  Skills across all profiles",str(total_cand_sk)],
        ["  Location preferences (total)",str(len(loc_prefs))],
        ["Total Job Postings",str(len(postings))],
        ["  Skills across all postings",str(total_comp_sk)],
        ["Total Applications",str(len(applications))],
        ["Total Matches",str(len(matches))],
        ["Total Swipes",str(len(swipes))],
        ["Avg profiles / candidate",f"{avg_pr:.2f}"],
        ["Avg certifications / candidate",f"{avg_ct:.2f}"],
        ["Avg postings / company",f"{avg_po:.2f}"],
        ["Avg skills / job profile",f"{avg_ps:.2f}"],
        ["Avg skills / job posting",f"{avg_jp:.2f}"],
    ],col_widths=[Inches(5.5),Inches(1.5)])
    doc.add_page_break()

    # ----- Section 2: Candidate Details ---------------------------------------
    _wdg_heading(doc,"2. Candidate Details")
    for cand in candidates:
        user=user_by_id.get(cand.user_id)
        c_res=resumes_by_cand.get(cand.id,[]); c_crt=certs_by_cand.get(cand.id,[])
        c_prof=profiles_by_cand.get(cand.id,[]); c_apps=apps_by_cand.get(cand.id,[])
        c_mat=matches_by_cand.get(cand.id,[])
        _wdg_heading(doc,f"Candidate #{cand.id}  -  {cand.name}",level=2)
        socials=[]
        if cand.linkedin_url:  socials.append(f"LinkedIn: {cand.linkedin_url}")
        if cand.github_url:    socials.append(f"GitHub: {cand.github_url}")
        if cand.portfolio_url: socials.append(f"Portfolio: {cand.portfolio_url}")
        _add_table(doc,["Field","Value"],[
            ["Name",cand.name],["Email",cand.email],
            ["Login Email",user.email if user else "N/A"],
            ["Phone",cand.phone],["Address",cand.residential_address],
            ["State / Zip",f"{cand.location_state}  {cand.location_zipcode}"],
            ["Social Links","\n".join(socials) if socials else "None"],
            ["Profile Complete",str(cand.profile_complete)],
            ["Resumes",str(len(c_res))],["Certifications",str(len(c_crt))],
            ["Job Profiles",str(len(c_prof))],
            ["Applications",str(len(c_apps))],["Matches",str(len(c_mat))],
        ],col_widths=[Inches(2.0),Inches(7.0)],header_bg="2E74B5")
        if c_crt:
            doc.add_paragraph()
            p=doc.add_paragraph(); r=p.add_run("Certifications")
            r.bold=True; r.font.size=Pt(10)
            _add_table(doc,["ID","Name","Issuer","Issued","Expires"],
                [[str(c.id),c.name,c.issuer or "N/A",c.issued_date or "N/A",c.expiry_date or "N/A"] for c in c_crt],
                col_widths=[Inches(0.4),Inches(3.5),Inches(2.0),Inches(1.0),Inches(1.0)],
                header_bg="375623")
        for jp in c_prof:
            jp_sk=skills_by_profile.get(jp.id,[]); jp_lc=locs_by_profile.get(jp.id,[])
            tech=[s for s in jp_sk if s.skill_category=="technical"]
            func=[s for s in jp_sk if s.skill_category=="functional"]
            soft=[s for s in jp_sk if s.skill_category=="soft"]
            doc.add_paragraph()
            p=doc.add_paragraph(); r=p.add_run(f"Job Profile #{jp.id}: {jp.profile_name}")
            r.bold=True; r.font.size=Pt(10); r.font.color.rgb=RGBColor(0x2E,0x74,0xB5)
            _add_table(doc,["Field","Value"],[
                ["Vendor / Product / Role",f"{jp.product_vendor or 'N/A'} / {jp.product_type or 'N/A'} / {jp.job_role or 'N/A'}"],
                ["Experience",f"{jp.years_of_experience} years"],
                ["Seniority",jp.seniority_level or "N/A"],
                ["Work Type",_v(jp.worktype)],["Employment",_v(jp.employment_type)],
                ["Salary Range",f"{_v(jp.salary_currency)}  {jp.salary_min:,.0f} - {jp.salary_max:,.0f}"],
                ["Visa Status",_v(jp.visa_status)],
                ["Location Prefs",", ".join(f"{l.city} {l.state}" for l in jp_lc) or "None"],
                ["Skills Total",f"{len(jp_sk)} (Tech: {len(tech)}, Func: {len(func)}, Soft: {len(soft)})"],
            ],col_widths=[Inches(2.0),Inches(7.0)],header_bg="833C00")
            if jp_sk:
                doc.add_paragraph()
                p2=doc.add_paragraph(); r2=p2.add_run(f"Skills for Profile #{jp.id}")
                r2.bold=True; r2.font.size=Pt(9)
                _add_table(doc,["Skill Name","Category","Proficiency (1-5)"],
                    [[s.skill_name,s.skill_category,str(s.proficiency_level)] for s in jp_sk],
                    col_widths=[Inches(4.5),Inches(1.5),Inches(1.0)],header_bg="375623")
        doc.add_paragraph()
    doc.add_page_break()

    # ----- Section 3: Company Details -----------------------------------------
    _wdg_heading(doc,"3. Company Details")
    for comp in companies:
        user=user_by_id.get(comp.user_id); c_posts=postings_by_comp.get(comp.id,[])
        _wdg_heading(doc,f"Company #{comp.id}  -  {comp.company_name}",level=2)
        _add_table(doc,["Field","Value"],[
            ["Company Name",comp.company_name],
            ["Contact Email",comp.company_email],
            ["Login Email",user.email if user else "N/A"],
            ["Employee Type",comp.employee_type],
            ["Department",comp.department or "N/A"],
            ["Location",comp.company_location or "N/A"],
            ["Phone",comp.phone_number or "N/A"],
            ["Website",comp.company_website or "N/A"],
            ["LinkedIn",comp.linkedin_profile or "N/A"],
            ["Hiring Focus",comp.hiring_focus or "N/A"],
            ["Profile Complete",str(comp.profile_complete)],
            ["Description",(comp.company_description or "N/A")[:300]],
            ["Total Postings",str(len(c_posts))],
        ],col_widths=[Inches(2.0),Inches(7.0)],header_bg="2E74B5")
        for jp in c_posts:
            jp_sk=pskills_by_posting.get(jp.id,[]); tech=[s for s in jp_sk if s.skill_category=="technical"]; soft=[s for s in jp_sk if s.skill_category=="soft"]
            doc.add_paragraph()
            p=doc.add_paragraph(); r=p.add_run(f"Job Posting #{jp.id}: {jp.job_title}")
            r.bold=True; r.font.size=Pt(10); r.font.color.rgb=RGBColor(0x83,0x3C,0x00)
            _add_table(doc,["Field","Value"],[
                ["Vendor / Product / Role",f"{jp.product_vendor} / {jp.product_type} / {jp.job_role}"],
                ["Seniority",jp.seniority_level],
                ["Work Type",_v(jp.worktype)],["Employment",_v(jp.employment_type)],
                ["Location",jp.location],["Start Date",jp.start_date],
                ["Salary",f"{_v(jp.salary_currency)}  {jp.salary_min:,.0f} - {jp.salary_max:,.0f}"],
                ["Status",f"{_v(jp.status)}  (Active: {jp.is_active})"],
                ["Skills Total",f"{len(jp_sk)} (Tech: {len(tech)}, Soft: {len(soft)})"],
            ],col_widths=[Inches(2.0),Inches(7.0)],header_bg="1F3864")
            if jp_sk:
                doc.add_paragraph()
                p2=doc.add_paragraph(); r2=p2.add_run(f"Required Skills for Posting #{jp.id}")
                r2.bold=True; r2.font.size=Pt(9)
                _add_table(doc,["Skill Name","Category","Rating (1-10)"],
                    [[s.skill_name,s.skill_category,str(s.rating)] for s in jp_sk],
                    col_widths=[Inches(4.5),Inches(1.5),Inches(1.0)],header_bg="375623")
            doc.add_paragraph()
    doc.add_page_break()

    # ----- Section 4: Count Tables --------------------------------------------
    _wdg_heading(doc,"4. Count Summary Tables")

    _wdg_heading(doc,"Table A: Candidates - Data Completeness",level=2)
    tableA=[]
    for cand in candidates:
        cp=profiles_by_cand.get(cand.id,[])
        tableA.append([str(cand.id),cand.name,
            str(len(resumes_by_cand.get(cand.id,[]))),
            str(len(certs_by_cand.get(cand.id,[]))),
            str(len(cp)),
            str(sum(len(skills_by_profile.get(p.id,[])) for p in cp)),
            str(sum(len(locs_by_profile.get(p.id,[])) for p in cp)),
            str(len(apps_by_cand.get(cand.id,[])))])
    _add_table(doc,["ID","Name","Resumes","Certs","Profiles","Prof Skills","Locations","Apps"],tableA,
        col_widths=[Inches(0.4),Inches(2.2),Inches(0.8),Inches(0.6),Inches(0.8),Inches(0.9),Inches(0.9),Inches(0.6)])

    doc.add_paragraph()
    _wdg_heading(doc,"Table B: Companies - Job Postings & Skills",level=2)
    tableB=[]
    for comp in companies:
        cp=postings_by_comp.get(comp.id,[])
        tableB.append([str(comp.id),comp.company_name,comp.employee_type,str(len(cp)),
            str(sum(len(pskills_by_posting.get(p.id,[])) for p in cp))])
    _add_table(doc,["ID","Company","Employee Type","Postings","Posting Skills"],tableB,
        col_widths=[Inches(0.4),Inches(2.8),Inches(1.8),Inches(0.9),Inches(1.2)])

    doc.add_paragraph()
    _wdg_heading(doc,"Table C: Job Profiles - Skills Breakdown",level=2)
    tableC=[]
    for jp in profiles:
        jp_sk=skills_by_profile.get(jp.id,[])
        tableC.append([str(jp.id),str(jp.candidate_id),jp.profile_name,
            str(sum(1 for s in jp_sk if s.skill_category=="technical")),
            str(sum(1 for s in jp_sk if s.skill_category=="functional")),
            str(sum(1 for s in jp_sk if s.skill_category=="soft")),
            str(len(jp_sk))])
    _add_table(doc,["PID","CID","Profile Name","Technical","Functional","Soft","Total"],tableC,
        col_widths=[Inches(0.4),Inches(0.4),Inches(4.5),Inches(0.9),Inches(1.0),Inches(0.6),Inches(0.6)])

    doc.add_page_break()
    _wdg_heading(doc,"Table D: Job Postings - Skills Breakdown",level=2)
    tableD=[]
    for jp in postings:
        jp_sk=pskills_by_posting.get(jp.id,[])
        tableD.append([str(jp.id),str(jp.company_id),jp.job_title,
            str(sum(1 for s in jp_sk if s.skill_category=="technical")),
            str(sum(1 for s in jp_sk if s.skill_category=="soft")),
            str(len(jp_sk))])
    _add_table(doc,["JID","CID","Job Title","Technical","Soft","Total"],tableD,
        col_widths=[Inches(0.4),Inches(0.4),Inches(5.5),Inches(0.9),Inches(0.7),Inches(0.7)])

    doc.add_paragraph()
    _wdg_heading(doc,"Table E: Social Links per Candidate",level=2)
    _add_table(doc,["CID","Name","LinkedIn","GitHub","Portfolio"],
        [[str(c.id),c.name,"Yes" if c.linkedin_url else "No",
          "Yes" if c.github_url else "No","Yes" if c.portfolio_url else "No"]
         for c in candidates],
        col_widths=[Inches(0.4),Inches(3.0),Inches(1.0),Inches(1.0),Inches(1.5)])

    doc.add_paragraph()
    _wdg_heading(doc,"Table F: Social Links per Company",level=2)
    _add_table(doc,["CID","Company","LinkedIn"],
        [[str(c.id),c.company_name,"Yes" if c.linkedin_profile else "No"] for c in companies],
        col_widths=[Inches(0.4),Inches(4.0),Inches(1.0)])

    doc.add_page_break()

    # ----- Section 5: Quick Stats ---------------------------------------------
    _wdg_heading(doc,"5. Quick Stats Snapshot  (Manager Summary)")
    _add_table(doc,["Metric","Value"],[
        ["Total users in system",str(len(users))],
        ["  - Candidates",str(len(candidates))],
        ["  - Companies",str(len(companies))],
        ["Avg job profiles per candidate",f"{avg_pr:.2f}"],
        ["Avg certifications per candidate",f"{avg_ct:.2f}"],
        ["Avg job postings per company",f"{avg_po:.2f}"],
        ["Avg skills per job profile",f"{avg_ps:.2f}"],
        ["Avg skills per job posting",f"{avg_jp:.2f}"],
        ["Total applications",str(len(applications))],
        ["Total matches",str(len(matches))],
        ["Total swipes",str(len(swipes))],
    ],col_widths=[Inches(5.5),Inches(1.5)])

    doc.save(output_path)
    return output_path

# =============================================================================
# MAIN
# =============================================================================

def main():
    ts=datetime.now().strftime("%Y-%m-%d_%H%M%S")
    output_file=os.path.join(SCRIPT_DIR, f"TalentGraph_Report_{ts}.docx")
    print(f"\nTalentGraph V2 Database Report - {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    print(f"Word report will be saved to: {output_file}\n")
    with Session(engine) as session:
        data=collect_data(session)
    print_report(data)
    if DOCX_AVAILABLE:
        saved=generate_word_report(data,output_file)
        if saved:
            print(f"\n[DONE] Word report saved: {saved}")
    else:
        print("\n[SKIP] Install python-docx to enable Word export.")

if __name__=="__main__":
    main()
